import markdown
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.template.loader import get_template, render_to_string
from django.utils.safestring import mark_safe
from markdownify.templatetags.markdownify import markdownify
from xhtml2pdf import pisa
import re

from RoleRadar import settings
from ai.charts import custom_analyze
from ai.classifires import TestClassifier
from ai.gemini import fetch_gemini_response
from ai.job_summery import JobSummery

from docx import Document
import io
import fitz
from typing import Union

from dashboard.models import ResumeAnalysis
from django.contrib import messages

import markdown
from docx import Document
from io import BytesIO
from bs4 import Tag, BeautifulSoup

model = TestClassifier()
job_summery = JobSummery()


@login_required
def dashboard(request):
    user = request.user
    analyses = ResumeAnalysis.objects.filter(user=request.user).order_by('-created_at')
    context = {
        'user': user,
        'analyses': analyses
    }
    return render(request, 'dashboard/dashboard.html', context)


@login_required
def resume_history(request):
    analyses = ResumeAnalysis.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'dashboard/includes/resume_history.html', {'analyses': analyses})


@login_required
def resume_detail(request, id):
    analysis = ResumeAnalysis.objects.get(id=id, user=request.user)
    return render(request, 'dashboard/resume_detail.html', {'analysis': analysis})


@login_required
def download_resume_analysis_pdf(request, id):
    analysis_obj = get_object_or_404(ResumeAnalysis, pk=id)
    context = {'analysis': analysis_obj}
    template_path = 'dashboard/resume_analysis_pdf_template.html'
    template = get_template(template_path)
    html_string = template.render(context)
    result_file = io.BytesIO()
    pisa_status = pisa.CreatePDF(
        src=html_string.encode('utf-8'),
        dest=result_file,
        encoding='utf-8'
    )

    if pisa_status.err:
        return HttpResponse(f"Error generating PDF: <pre>{html_string}</pre> <pre>{pisa_status.err}</pre>")
    response = HttpResponse(result_file.getvalue(), content_type='application/pdf')
    response[
        'Content-Disposition'] = f'attachment; filename="resume_analysis_{analysis_obj.id}_{analysis_obj.user.username}.pdf" '

    result_file.close()
    return response


def getPDFText(file_stream: bytes) -> Union[str, None]:
    page_texts = []
    doc = None
    try:
        doc = fitz.open(stream=file_stream, filetype="pdf")
        if not doc or doc.page_count == 0:
            raise ValueError("Uploaded PDF is empty or unreadable.")
        for page in doc:
            page_texts.append(page.get_text("text"))

        if not page_texts:
            return ""
        return "\n".join(page_texts)
    except Exception as e:
        print(f"Error opening or reading PDF stream: {type(e).__name__} - {e}")
        return None
    finally:
        if doc:
            doc.close()


def getDocxText(uploaded_file_object) -> Union[str, None]:
    try:
        uploaded_file_object.seek(0)

        file_stream = io.BytesIO(uploaded_file_object.read())

        document = Document(file_stream)
        if not document.paragraphs:
            raise ValueError("Uploaded DOCX is empty or unreadable.")
        full_text = [para.text for para in document.paragraphs]
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading .docx file: {type(e).__name__} - {e}")
        return None


@login_required
def key_skills(request):
    user = request.user
    if request.method == 'POST':
        description = request.POST.get('description')
        uploaded_file = request.FILES.get('resume_file')

        if uploaded_file and uploaded_file.size > settings.MAX_FILE_SIZE_MB * 1024 * 1024:
            return render(request, 'dashboard/key_skills.html', {
                'user': user,
                'error': "File size must be less than 5MB."
            })

        if not description and uploaded_file:
            if uploaded_file.name.endswith('.pdf'):
                description = getPDFText(uploaded_file.read())
            elif uploaded_file.name.endswith('.docx') or uploaded_file.name.endswith('.doc'):
                description = getDocxText(uploaded_file)
            else:
                description = None

        results, top_3_items = model.predict(description)
        job_titles, job_summery_text = job_summery.get_job_summery(top_3_items[0])
        text = f"Review the resume with the job description. \n resume: \n {description} \n job description:  \n {job_summery_text} \n Tell me what the shortcomings of your resume are and what you need to improve. Tell me what the shortcomings are in your resume and what you need to improve, and finally, give me a roadmap for improvement. "
        prompt_resume_analysis = f"""
        Analyze the following resume with respect to the given job description.
        Resume:
        {description}

        Job Description:
        {job_summery_text}

        Identify the weaknesses and missing elements in this resume for this job.
        """
        prompt_resume_roadmap = f"""
        Based on the analysis of this resume {description} and job description {job_summery_text}, provide a detailed roadmap to improve the resume.
        Be specific about what sections to add, modify, or remove.
        """

        prompt_skill_roadmap = f"""
        Given this resume and the following job description, provide a skill development roadmap.
        Mention which technical and soft skills the candidate should learn or improve, and in what order, with approximate timelines.

        Resume:
        {description}

        Job Description:
        {job_summery_text}
        """
        suggestion_analysis = fetch_gemini_response(prompt_resume_analysis)
        suggestion_resume = fetch_gemini_response(prompt_resume_roadmap)
        suggestion_skills = fetch_gemini_response(prompt_skill_roadmap)

        suggestion = fetch_gemini_response(prompt_content=text)

        ResumeAnalysis.objects.create(
            user=user,
            original_resume=description,
            predictions=results,
            top_job=top_3_items[0],
            job_summary=job_summery_text,
            # suggestion=suggestion
        )

        print(job_titles[0])
        analysis_results = custom_analyze(job_title=job_titles[0][0])
        context = {
            'user': user,
            'results': results,
            'top_3_items': top_3_items,
            'job_summery_text': job_summery_text,
            'job_summery_text_only': markdown_to_text(job_summery_text),
            'suggestion': suggestion,
            'job_title': job_titles[0][0],
            'analysis_results': analysis_results,
            'suggestion_analysis': suggestion_analysis,
            'suggestion_analysis_text': markdown_to_text(suggestion_analysis),
            'suggestion_resume': suggestion_resume,
            'suggestion_resume_text': markdown_to_text(suggestion_resume),
            'suggestion_skills': suggestion_skills,
            'suggestion_skills_text': markdown_to_text(suggestion_skills)

        }
        return render(request, 'dashboard/key_skills.html', context)
    else:
        context = {
            'user': user
        }
        return render(request, 'dashboard/key_skills.html', context)


def markdown_to_text(text):
    html = markdown.markdown(text)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text()


@login_required
def roadmap(request):
    user = request.user
    if request.method == 'POST':
        description = request.POST.get('description')
        text = f"Give a roadmap on the following topic. \n {description}"
        message = fetch_gemini_response(text)

        context = {
            'user': user,
            'message': message,
            'text_message': markdown_to_text(message),
        }
        return render(request, 'dashboard/roadmap.html', context)
    else:
        context = {
            'user': user
        }
        return render(request, 'dashboard/roadmap.html', context)


@login_required
def resume(request):
    user = request.user
    context = {'user': user}

    if request.method == 'POST':

        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        email = request.POST.get('email', '').strip()
        phone_number = request.POST.get('phone_number', '').strip()
        linkedin_url = request.POST.get('linkedin_url', '').strip()
        portfolio_url = request.POST.get('portfolio_url', '').strip()

        skills_raw = request.POST.get('skills', '')
        description_text = request.POST.get('description', '').strip()

        prompt_parts = [
            "Please generate a professional resume draft based on the following details."
            "or anything out of resume! "
            "Organize the information into standard resume sections (e.g., Contact Information, Summary, Skills, "
            "Professional Links). "
            "Ensure the tone is professional and the content is well-written."
        ]

        personal_details_section = []
        if first_name or last_name:
            personal_details_section.append(f"Full Name: {first_name} {last_name}".strip())
        if email:
            personal_details_section.append(f"Email: {email}")
        if phone_number:
            personal_details_section.append(f"Phone: {phone_number}")

        if personal_details_section:
            prompt_parts.append("\n[Contact Information]")
            prompt_parts.extend(personal_details_section)

        links_section = []
        if linkedin_url:
            links_section.append(f"LinkedIn: {linkedin_url}")
        if portfolio_url:
            links_section.append(f"Portfolio/Website: {portfolio_url}")

        if links_section:
            prompt_parts.append("\n[Professional Links]")
            prompt_parts.extend(links_section)

        if description_text:
            prompt_parts.append(f"\n[Summary/About Me]\n{description_text}")

        if skills_raw:
            skills_list = [s.strip() for s in skills_raw.split(',') if s.strip()]
            if skills_list:
                prompt_parts.append("\n[Skills]")
                for skill in skills_list:
                    prompt_parts.append(f"- {skill}")

        final_prompt = "\n".join(prompt_parts)

        print(f"DEBUG: Prompt sent to Gemini:\n{final_prompt}")

        message_from_gemini = fetch_gemini_response(final_prompt)
        request.session['generated_resume_markdown'] = message_from_gemini
        context['message'] = message_from_gemini
        context['form_data'] = request.POST

        return render(request, 'dashboard/resume.html', context)

    return render(request, 'dashboard/resume.html', context)


@login_required
def advice(request):
    user = request.user
    if request.method == 'POST':
        text = request.POST.get('text')
        message = fetch_gemini_response(text)

        html = markdownify(message)
        code_block_pattern = re.compile(r'```(\w+)?\n(.*?)```', re.DOTALL)

        def replace_code_block(match):
            language = match.group(1) or ''
            code_content = match.group(2)
            escaped_code = (
                code_content.replace('&', '&amp;')
                    .replace('<', '&lt;')
                    .replace('>', '&gt;')
            )
            return f'<pre><code class="{language}">{escaped_code}</code></pre>'

        final_html = re.sub(code_block_pattern, replace_code_block, html)

        return HttpResponse(mark_safe(final_html))

    else:
        context = {
            'user': user,
            'context': request.GET.get('context')
        }
        return render(request, 'dashboard/advice.html', context)


@login_required
def resume_summerize(request):
    user = request.user

    if request.method == 'POST':
        description = request.POST.get('description')
        uploaded_file = request.FILES.get('resume_file')

        if uploaded_file and uploaded_file.size > settings.MAX_FILE_SIZE_MB * 1024 * 1024:
            return render(request, 'dashboard/resume_summerize.html', {
                'user': user,
                'error': "File size must be less than 5MB."
            })

        if not description and uploaded_file:
            if uploaded_file.name.endswith('.pdf'):
                description = getPDFText(uploaded_file.read())
            elif uploaded_file.name.endswith('.docx') or uploaded_file.name.endswith('.doc'):
                description = getDocxText(uploaded_file)
            else:
                description = None

        if not description:
            return render(request, 'dashboard/resume_summerize.html', {
                'user': user,
                'error': "Please provide a description or upload a valid resume file."
            })

        text = f"Summarize this resume:\n{description}"
        message = fetch_gemini_response(text)

        context = {
            'user': user,
            'message': message
        }
        return render(request, 'dashboard/resume_summerize.html', context)

    else:
        context = {
            'user': user
        }
        return render(request, 'dashboard/resume_summerize.html', context)


@login_required
def delete_resume_analysis_view(request, id):
    analysis_to_delete = get_object_or_404(ResumeAnalysis, id=id,
                                           user=request.user)

    try:
        analysis_to_delete.delete()
        messages.success(request, f"Analysis for '{analysis_to_delete.top_job}' was successfully deleted.")
    except Exception as e:
        messages.error(request, f"Error deleting analysis: {e}")
        print(f"Error deleting analysis (ID: {id}): {e}")

    return redirect('dashboard')


@login_required
def download_generated_resume_pdf(request):
    markdown_content = request.session.get('generated_resume_markdown', None)

    if not markdown_content:
        return HttpResponse("No resume content found to generate PDF. Please generate a resume first.", status=404)

    html_content_raw = markdown.markdown(markdown_content, extensions=['extra', 'nl2br', 'sane_lists'])

    context_for_pdf = {'content': html_content_raw}

    final_html_for_pdf = render_to_string('dashboard/pdf_resume_from_markdown.html', context_for_pdf)

    result_file = io.BytesIO()
    pisa_status = pisa.CreatePDF(
        src=final_html_for_pdf.encode('utf-8'),
        dest=result_file,
        encoding='utf-8'
    )

    if pisa_status.err:
        return HttpResponse(f"Error generating PDF: {pisa_status.err}")

    response = HttpResponse(result_file.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="generated_resume.pdf"'
    result_file.close()
    return response


def markdown_to_word(markdown_text, filename='resume'):
    html_content = markdown.markdown(markdown_text)

    document = Document()

    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    for element in soup.contents:
        if isinstance(element, Tag):
            if element.name == 'h1':
                document.add_heading(element.text, level=1)
            elif element.name == 'h2':
                document.add_heading(element.text, level=2)
            elif element.name == 'h3':
                document.add_heading(element.text, level=3)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    document.add_paragraph(li.text, style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    document.add_paragraph(li.text, style='List Number')
            elif element.name == 'p':
                document.add_paragraph(element.text)
            else:
                document.add_paragraph(element.text)

    byte_io = BytesIO()
    document.save(byte_io)
    byte_io.seek(0)

    class DocxFile:
        def __init__(self, bytes_data):
            self.bytes = bytes_data

    return DocxFile(byte_io.read())


@login_required
def download_generated_resume_word(request):
    markdown_content = request.session.get('generated_resume_markdown', None)

    if not markdown_content:
        return HttpResponse(
            "No resume content found to generate Word document. Please generate a resume first.",
            status=404
        )

    try:

        file_name = f"{request.user.first_name}_resume"
        converter = markdown_to_word(markdown_content, file_name)
        docx_bytes = converter.bytes

        response = HttpResponse(
            docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{file_name}.docx"'
        return response

    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return HttpResponse(f"Error generating Word document: {e}", status=500)
